# -*- coding: utf-8 -*-
"""重建简历 docx:参考网络优质简历结构(个人总结→教育→技能→项目),内容全部基于代码事实"""
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PATH = '_new_resume.docx'
d = docx.Document()

for s in d.sections:
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

FONT = '微软雅黑'


def set_font(run, size, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._r.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def para(segments, size=10.5, align=None, space_before=0, space_after=2, line=1.3):
    p = d.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align:
        p.alignment = align
    for text, bold in segments:
        r = p.add_run(text)
        set_font(r, size, bold)
    return p


def section(title):
    para([(title, True)], size=13, space_before=8, space_after=4)


# ---------- 头部:姓名 + 意向 + 联系方式 ----------
para([('陈垲霖', True)], size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para([('前端开发（实习 / 初级） ｜ 深圳 ｜ 可随时到岗', False)], size=10.5,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
para([('18689469258 ｜ 1300949382@qq.com ｜ https://ckl1300949382.github.io/-vue3-/', False)], size=10.5,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# ---------- 个人总结 ----------
section('个人总结')
para([('广州科技职业技术大学计算机应用工程 2027 届。在校独立完成并上线一个 Vue 3 + TypeScript 后台管理系统，', False)], space_after=0)
para([('走通需求分析、架构搭建、开发调试、部署上线全流程；注重组件拆分与逻辑复用，习惯借助 AI 工具辅助调试提效。', False)], space_after=4)

# ---------- 教育经历 ----------
section('教育经历')
para([('2023.09 - 2027.06  广州科技职业技术大学', True)], size=11, space_after=1)
para([('计算机应用工程 ｜ 本科', False)], space_after=4)

# ---------- 专业技能 ----------
section('专业技能')
skills = [
    'HTML / CSS：熟悉语义化标签、Flex / Grid 布局，能完成响应式页面',
    'JavaScript / TypeScript：熟练 ES6+ 语法，理解事件循环与异步（Promise / async-await）；项目中使用 interface、泛型约束数据类型',
    'Vue 3：熟悉 Composition API（script setup）、Vue Router（导航守卫、路由懒加载）、Pinia（状态持久化）',
    'Axios：封装请求/响应拦截器，统一 Token 注入与错误提示，401 自动退出登录',
    '组件库 / 可视化：使用 Element Plus 开发后台页面，ECharts 图表开发',
    '工程工具：日常使用 Git 管理代码，理解 Vite 构建流程，可独立完成项目搭建与部署',
]
for s in skills:
    para([('- ', False), (s, False)])

# ---------- 项目经历 ----------
section('项目经历')
para([('【项目一】可视化用户管理平台（个人项目）｜ 2026.05 - 2026.08', True)], size=11, space_after=1)
para([('技术栈：Vue 3 + TypeScript + Composition API + Pinia + Vue Router + Element Plus + Axios + ECharts + Express', False)], size=10.5, space_after=3)

projects = [
    '独立开发并上线：从需求分析到部署，按 api / store / router / views 分层组织代码；自建 types/ 类型契约（BizResult<T> 统一响应、PageData<T> 分页结构），13 个接口调用全程类型约束',
    '搭建登录鉴权：Pinia 持久化登录态，刷新页面不丢失；路由全局守卫按 route.meta 控制访问权限（需登录、需管理员），未登录或越权自动提示并跳转',
    '封装请求层：axios 拦截器统一注入 Token、集中弹窗报错，401 自动清除登录态并跳转登录页，页面无需重复处理异常',
    '拆分复用：用户管理页拆为 UserTable（表格）、UserFormDialog（新增/编辑共用弹窗）、UserManage（组装逻辑）三组件，新增与编辑复用同一表单；抽离 5 个组合式函数（防抖搜索、分页、列表逻辑）',
    '数据可视化：集成 4 个 ECharts 图表（注册趋势、角色分布、状态分布、月度注册），组件销毁时释放实例避免内存泄漏',
    '前后端联调：配合 Express + JWT 后端（13 个接口），独立编写 761 行《API 接口文档》规范联调口径',
    '部署上线：配置 hash 路由与静态构建，已部署 GitHub Pages 可在线访问',
]
for desc in projects:
    para([('- ', False), (desc, False)])

d.save(PATH)
print('OK')
